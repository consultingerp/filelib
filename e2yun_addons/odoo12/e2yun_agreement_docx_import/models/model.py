# -*- coding: utf-8 -*-

import base64
from odoo import api, fields, models, tools, _
from docx import Document
from docx.shared import Pt,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os, sys
import platform
class DocxImport(models.TransientModel):
    _name = "agreement.docx.import"
    _description = "agreement docx Import"

    agreement_id = fields.Many2one(
        'agreement', string='Agreement', ondelete='restrict', required=True)
    name = fields.Char('file path')
    data = fields.Binary('File')
    filename = fields.Char('File Name')

    @api.multi
    def import_docx(self):
        this = self[0]
        full_path = this.name
        doc = Document(full_path)

        recital = self.env['agreement.recital']    #叙述
        clause = self.env['agreement.clause']      #条款
        section = self.env['agreement.section']    #章节
        appendix = self.env['agreement.appendix']  #附录




        paragraphs = doc.paragraphs
       #paragraphs[6].text = "小Z同学："
        #paragraphs[6].add('some striked', strike=True)
        val = {}
        val['testxml'] =paragraphs[6]._element.xml
        wdel = '<w: delw: id="0"w: author="pactera"w: date="2019-12-02T15:56:00Z"><w: rw: rsidDel="00C410A9"><w: rPr><w: rFontsw: ascii="华文中宋"w: eastAsia="华文中宋"w: hAnsi="华文中宋"w: hint="eastAsia"/><w: szw: val="32"/><w: szCsw: val="32"/></w: rPr><w: delText>北京海辉高科软件有限公司</w: delText></w: r></w: del>'
        #num = xml.find('合同号')
        #str(doc.element.xml).replace(str(doc.element.xml)[num], "徐初秋")
        #paragraphs[6]._element.xml = val['testxml']
        #doc.Comments.Add(Range=doc.paragraphs[6].Range, Text='注释') #其中我的i写的是遍历整个段落，然后匹配出需要加注释的地方

        doc.save('D:\\demo1207.docx')
        vals=[]
       # 每一段的内容
        i=0
        for para in doc.paragraphs:
            print(i)
            val = {}
            val['no'] =i
            if para.text=='实施服务协议':
                print(1)
            val['text']=para.text
            val['style_font'] = para.style.font.name
            val['style_name'] = para.style.name
            val['style'] = para.style
            val['font'] = para.style.font
            vals.append(val)
            i=i+1

        document = Document() # 初始化文档
        for val in vals:
            print(val)
            if val['style_name'] and val['style_name']=='文档标题':
                d_title = document.add_heading(val['text'], 0)
                d_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                d_title.style=val['style']
                #d_title.style.font.size=28
                #d_title.bold = True

            else:
              p=document.add_paragraph(val['text'])
              p.style=val['style']

        document.save('D:\\demo1203.docx')

        return True



class AgreementDownloadDoc(models.Model):
    _name = "agreement.download.doc"

    file_path = fields.Char('File Path')
    data = fields.Binary('File')
    filename = fields.Char('File Name')

    #默认样式设置
    def chg_font(self,obj, fontname='微软雅黑', size=None):
         ## 设置字体函数
        obj.font.name = fontname
        #obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
        if size and isinstance(size, Pt):
          obj.font.size = size

    def add_text_default(self,document,text,bold,size,sytle_name,font_name):
        p = document.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.style.font.size = size
        r.style.name =sytle_name
        r.style.font.name=font_name

    def wb_path(self,name):
        platform_ = platform.system()
        if platform_ == "Windows":
            wb_path = "" + str(sys.path[0]) + "/"+str(name)+"Temp.docx"
        else:
            wb_path = "/tmp/"+str(name)+"Temp.docx"
        return  wb_path

    def download_doc(self):
     try:
        wb_path = self.wb_path('download_doc')

        agreement_id=self._context['active_id']

        attachment = self.env['ir.attachment'] # 附件

        agreement = self.env['agreement']

        agreementData = agreement.browse(agreement_id)

        agreement_word_data = self.env['agreement.word.data']  # wordData

        clauseListData = agreement_word_data.search([('agreement_id', '=', agreement_id),
                                                     ('detail','=',False),('the_editor','=',True)])

        master_word_id=0

        if clauseListData:
            master_word_id=clauseListData[0].master_word_id
        else:
            clauseListData = agreement_word_data.search(
                [('agreement_id', '=', agreement_id), ('detail', '=', False)])
            master_word_id = clauseListData[0].master_word_id


        word_data=self.env['ir.http'].binary_content(
            xmlid=None, model='ir.attachment', id=master_word_id ,field='datas')

        f = open(wb_path, r"wb")
        f.write(base64.decodestring(word_data[2]))
        f.close()
        doc = Document(wb_path)

        try:
              for clauseObj in clauseListData:  # 条款
                  para = doc.paragraphs[clauseObj.sequence]
                  if clauseObj.the_editor==True:
                        if clauseObj.edit_type=='replace':
                            para.text=para.text.replace(clauseObj.old_text,clauseObj.new_text,1)
                        else:
                            para.text=clauseObj.new_text
              doc.save(wb_path)
        except BaseException as e:
            if os.path.exists(wb_path):
                os.remove(wb_path)


        file = open(wb_path, "rb")
        attachment.search([('res_model', '=', 'agreement.download.doc'),('res_id', '=',agreement_id),
                           ('res_name', '=','download_doc')]).unlink()  #删除无效附件
        attachmentObj=attachment.create({
            'name': agreementData.name,
            'datas': base64.encodestring(file.read()),
            'datas_fname': agreementData.name+".docx",
            'res_model': 'agreement.download.doc',
            'res_id': agreement_id,
            'res_name':'download_doc'
        })
        file.close()

        if os.path.exists(wb_path):
            os.remove(wb_path)

        return {
            'type': 'ir.actions.act_url',
            'target': 'new',
            'url': 'web/content/%s?download=true' % (attachmentObj.id),
        }

     except BaseException as e:
         print(e)



    def import_recital(self):
        wb_path=self.wb_path('recital')
        f = open(wb_path, r"wb")
        datass = base64.decodestring(self.data)
        f.write(datass)
        f.close()
        full_path = wb_path
        doc = Document(full_path)

        agreement_id = self._context['active_id']  # 读取当前合同ID

        agreement_recital_obj = self.env['agreement.recital']  # 附录对象

        attachment = self.env['ir.attachment']  # 附件

        # 记录当前 概述 母版 word

        file = open(full_path, "rb")
        attachment.search([('res_model', '=', 'agreement.download.doc'),
                           ('res_id', '=',self.id),('res_name', '=','agreement.recital')]).unlink()

        attachment_id = attachment.create({
            'name': '',
            'datas': base64.encodestring(file.read()),
            'datas_fname': "",
            'res_model': 'agreement.download.doc',
            'res_name': 'agreement.recital',
            'res_id': self.id
        })
        file.close()

        vals = []
        # 每一段的内容
        i = 0
        # 利用下标遍历段落
        all_content=""
        val = {}
        val['agreement_id'] = agreement_id
        val['sequence'] = 0
        val['master_word_id'] = attachment_id.id
        for i in range(len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if para.text or para.text != "":
                if para.style.font.size:
                    font_size = para.style.font.size
                else:
                    font_size=11
                text = "<p style= 'font-size: " + str(font_size) + "px; "
                if str(para.alignment)=="CENTER (1)":
                    text=text+" text-align:center; "

                runs_len=len(para.runs)-1
                if para.runs[runs_len].bold and para.runs[runs_len].bold==True:
                    text = text + " font-weight:bold; "

                text = text+" ' >"+para.text+"</p>"

                all_content = all_content+text

        # html = PyDocX.to_html(r"" + full_path)
        # val['name'] = self.filename
        # val['content'] = html
        # vals.append(val)

        val['name'] = "概述"
        val['content']=all_content
        vals.append(val)
        agreement_recital_obj.search([('agreement_id', '=', agreement_id)]).unlink()
        agreement_recital_obj.create(vals)

        if os.path.exists(wb_path):
            os.remove(wb_path)
        return True

    def import_sections_articles(self):

        wb_path = self.wb_path('sections_articles')
        f = open(wb_path, r"wb")
        datass = base64.decodestring(self.data)
        f.write(datass)
        f.close()
        full_path = wb_path
        doc = Document(full_path)

        agreement_id = self._context['active_id']  # 读取当前合同ID

        agreement_section_obj = self.env['agreement.section']  # 章节条款

        attachment = self.env['ir.attachment']  # 附件

        # 记录当前 章节条款 母版 word

        file = open(full_path, "rb")

        attachment.search([('res_model', '=', 'agreement.download.doc'),
                           ('res_id', '=', self.id), ('res_name', '=', 'agreement.section')]).unlink()

        attachment_id = attachment.create({
            'name': '',
            'datas': base64.encodestring(file.read()),
            'datas_fname': "",
            'res_model': 'agreement.download.doc',
            'res_name': 'agreement.section',
            'res_id': self.id
        })
        file.close()

        vals = []
        val = {}
        val['agreement_id'] = agreement_id
        val['sequence'] = 0
        val['master_word_id'] = attachment_id.id
        all_content=""
        for i in range(len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if para.text or para.text != "":
                if para.style.font.size:
                    font_size = para.style.font.size
                else:
                    font_size = 11
                text = "<p style= 'font-size: " + str(font_size) + "px; "
                if str(para.alignment) == "CENTER (1)":
                    text = text + " text-align:center; "

                runs_len = len(para.runs) - 1
                if para.runs[runs_len].bold and para.runs[runs_len].bold == True:
                    text = text + " font-weight:bold; "

                text = text + " ' >" + para.text + "</p>"

                all_content = all_content + text

        # html = PyDocX.to_html(r""+full_path)
        # val['name'] = self.filename
        # val['content'] = html

        val['name'] = "章节-条款"
        val['content'] = all_content

        vals.append(val)
        agreement_section_obj.search([('agreement_id', '=', agreement_id)]).unlink()
        agreement_section_obj.create(vals)

        if os.path.exists(wb_path):
            os.remove(wb_path)
        return True

    def import_appendix(self):
        wb_path = self.wb_path('appendix')
        f = open(wb_path, r"wb")
        datass = base64.decodestring(self.data)
        f.write(datass)
        f.close()
        full_path = wb_path
        doc = Document(full_path)

        agreement_id = self._context['active_id']  # 读取当前合同ID

        agreement_appendix_obj = self.env['agreement.appendix']  # 附录对象

        attachment = self.env['ir.attachment']  # 附件

        # 记录当前 概述 母版 word

        file = open(full_path, "rb")
        attachment.search([('res_model', '=', 'agreement.download.doc'),
                           ('res_id', '=', self.id), ('res_name', '=', 'agreement.appendix')]).unlink()

        attachment_id = attachment.create({
            'name': '',
            'datas': base64.encodestring(file.read()),
            'datas_fname': "",
            'res_model': 'agreement.download.doc',
            'res_name': 'agreement.appendix',
            'res_id': self.id
        })
        file.close()

        vals = []
        # 每一段的内容
        i = 0
        # 利用下标遍历段落
        all_content = ""
        val = {}
        val['agreement_id'] = agreement_id
        val['sequence'] = 0
        val['master_word_id'] = attachment_id.id
        for i in range(len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if para.text or para.text != "":
                if para.style.font.size:
                    font_size = para.style.font.size
                else:
                    font_size = 11
                text = "<p style= 'font-size: " + str(font_size) + "px; "
                if str(para.alignment) == "CENTER (1)":
                    text = text + " text-align:center; "

                runs_len = len(para.runs) - 1
                if para.runs[runs_len].bold and para.runs[runs_len].bold == True:
                    text = text + " font-weight:bold; "

                text = text + " ' >" + para.text + "</p>"

                all_content = all_content + text

        val['title'] = "附录"
        val['name'] = "附录"
        val['content'] = all_content
        vals.append(val)
        agreement_appendix_obj.search([('agreement_id', '=', agreement_id)]).unlink()
        agreement_appendix_obj.create(vals)

        if os.path.exists(wb_path):
            os.remove(wb_path)
        return True

    def Import_doc(self):
        # 分类型导入
        if 'import_type' in self._context.keys():
            import_type = self._context['import_type']
            if import_type == 'import_recital':
                self.import_recital()
            elif import_type == 'import_sections_articles':
                self.import_sections_articles()
            elif import_type == 'import_appendix':
                self.import_appendix()
            return True

        platform_ = platform.system()
        if platform_ == "Windows":
            wb_path = ""+str(sys.path[0]) +"/agreementTemp.docx"
        else:
            wb_path = "/tmp/agreementTemp.docx"

        f = open(wb_path, r"wb")
        datass = base64.decodestring(self.data)
        f.write(datass)
        f.close()
        full_path = wb_path

        doc = Document(full_path)

        agreement_id=self._context['active_id']   #读取当前合同ID
        agreement_word_data = self.env['agreement.word.data']  # wordData
        attachment = self.env['ir.attachment']  # 附件

        #记录当前合同的母版word
        file = open(full_path, "rb")
        attachment.search([('res_model', '=', 'agreement.download.doc'),
                           ('res_id', '=',agreement_id),
                           ('res_name', '=', 'Import_doc')]).unlink()
        attachment_id=attachment.create({
            'name': '',
            'datas': base64.encodestring(file.read()),
            'datas_fname': "",
            'res_model': 'agreement.download.doc',
            'res_id': agreement_id,
            'res_name':'Import_doc'
        })
        file.close()

        vals = []
        # 每一段的内容
        i = 0
        # 利用下标遍历段落
        for i in range(len(doc.paragraphs)):
            para = doc.paragraphs[i]
            val = {}
            val['agreement_id']=agreement_id
            val['sequence'] = i
            if para.text or  para.text!="":
                font_size=16
                val['alignment'] = str(para.alignment)  # 0,1,2 分别对应左对齐、居中、右对齐
                val['font_name'] = para.style.font.name
                val['font_size'] = font_size
                val['content'] = para.text
            val['master_word_id']=attachment_id.id
            vals.append(val)
        agreement_word_data.search([('agreement_id', '=', agreement_id)]).unlink()
        agreement_word_data.create(vals)
        if os.path.exists(wb_path):
            os.remove(wb_path)
        return True