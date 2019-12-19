# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.

import base64
from odoo import api, fields, models, tools, _
from docx import Document
import copy
import os, sys
import platform
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import difflib
from datetime import datetime, timedelta
class Agreement(models.Model):  #合同
    _inherit = "agreement"

    def wordEdit(self):

        return {
            'type': 'ir.actions.act_url',
            'target': 'new',
            'url': '/agreement/wordEdit/%(id)s' %  {'id': self.id},

        }
        # print(111)
        # agreement_word_data = self.env['agreement.word.data']  # wordData
        #
        # clauseListData = agreement_word_data.search(
        #     [('agreement_id', '=', self.id), ('detail', '=', True)])
        #
        #
        # return {
        #     'type': 'ir.actions.act_window',
        #     'res_model':'agreement.word.data',
        #     'view_type': 'form',
        #     'view_mode': 'form',
        #     'target': 'new',
        #     'res_id': clauseListData[0].id,
        #     'context': dict(self._context),
        # }

    def preview(self):

        return {
            'type': 'ir.actions.act_url',
            'target': 'current',
            'url': '/agreement/preview/%(id)s' % {'id': self.id},

        }

    def wb_path(self,name):
        platform_ = platform.system()
        if platform_ == "Windows":
            wb_path = "" + str(sys.path[0]) + "/"+str(name)+"Temp.docx"
        else:
            wb_path = "/tmp/"+str(name)+"Temp.docx"
        return  wb_path

    def download_word(self):

        try:
            wb_path = self.wb_path('download_word')

            agreement_id = self.id

            attachment = self.env['ir.attachment']  # 附件

            agreement_obj = self.env['agreement']

            agreementData = agreement_obj.browse(agreement_id)

            master_word_id = agreementData.recital_ids[0].master_word_id
            data_recital = self.env['ir.http'].binary_content(
                xmlid=None, model='ir.attachment', id=master_word_id, field='datas')

            master_word_id = agreementData.sections_ids[0].master_word_id
            data_sections = self.env['ir.http'].binary_content(
                xmlid=None, model='ir.attachment', id=master_word_id, field='datas')

            master_word_id = agreementData.appendix_ids[0].master_word_id
            data_appendix = self.env['ir.http'].binary_content(
                xmlid=None, model='ir.attachment', id=master_word_id, field='datas')
            doc = Document()
            i=0;
            while i<=3  :
              i=i+1
              if i==1:
                word_data=data_recital[2]
              elif i==2:
                word_data =data_sections[2]
              elif i==3:
                word_data =data_appendix[2]

              f = open(wb_path, r"wb")
              f.write(base64.decodestring(word_data))
              f.close()
              word_temp=Document(wb_path)
              for j in range(len(word_temp.paragraphs)):
                  paragraph = word_temp.paragraphs[j]
                  if not paragraph.text or paragraph.text =="" :
                      continue
                  #list_runs = copy.deepcopy(paragraph.runs)

                  p = doc.add_paragraph(paragraph.text)

                  if i==1 and j==5:
                      p.text = p.text.replace("XXXXXX", "中国远大", 1)
                      from docx.shared import RGBColor
                      test = p.runs[0]
                      font = test.font
                      font.color.rgb = RGBColor(220, 20, 60)


                  p._element = paragraph._element
                  p._p== paragraph._p
                  p._parent=paragraph._parent
                  p.alignment = paragraph.alignment
                  #p.paragraph_format = copy.deepcopy(paragraph.paragraph_format)
                  #p.paragraph_format.element=paragraph.paragraph_format.element
                  p.paragraph_format.alignment = paragraph.paragraph_format.alignment
                  print(paragraph.paragraph_format.space_before)
                  # p.paragraph_format.space_before =Pt(3)  # 上行间距
                  # p.paragraph_format.space_after ==Pt(12)# 下行间距
                  # p.paragraph_format.line_spacing =Pt(9)  # 行距


                  #p.part = paragraph.part
                 # p.text = paragraph.text
                  p.add_run=paragraph.runs
                  p.style = paragraph.style

                  #print(p)


                  # 文字替换
                  # for run in list_runs:
                  #     # for name in dic:
                  #     #     print(name)
                  #     #     if name in run.text:
                  #     #         value = dic[name]
                  #     #         run.text = run.text.replace(name, str(value))
                  #     # 段落样式的复制
                  #    print("")

              if os.path.exists(wb_path):
                  os.remove(wb_path)

            # section = doc.Sections(1)
            # footers = section.Footers  # a HeadersFooters collection object
            # default_footer = footers('wdHeaderFooterPrimary')
            # default_footer.Range.Text = "Footer text"

            doc.save(wb_path)


            file = open(wb_path, "rb")
            attachment.search(
                [('res_model', '=', 'agreement.download.doc'), ('res_id', '=', self.id),('res_name', '=','download_word')]).unlink()  # 删除无效附件

            agreementData.write_date = datetime.now()
            version=str(agreementData.version)+"."+str(agreementData.revision)
            attachmentObj = attachment.create({
                'name': agreementData.name,
                'datas': base64.encodestring(file.read()),
                'datas_fname': agreementData.name + "版本"+version+".docx",
                'res_model': 'agreement.download.doc',
                'res_id': self.id,
                'res_name':'download_word'
            })

            file.close()
            if os.path.exists(wb_path):
                os.remove(wb_path)

            return {
                'type': 'ir.actions.act_url',
                'target': 'new',
                'url': 'web/content/%s?download=true' % (attachmentObj.id),
            }

        except BaseException as e:
            print(e)


    def import_recital(self):
        wizard = self.env.ref(
            'e2yun_agreement_docx_import.view_action_agreement_import_doc')

        return {
            'name': '导入概述',
            'type': 'ir.actions.act_window',
            'view_type': 'form',
            'view_mode': 'form',
            'res_model': 'agreement.download.doc',
            'views': [(wizard.id, 'form')],
            'view_id': wizard.id,
            'target': 'new',
            'context': {
                'active_id': self.id,
                'import_type':'import_recital'
            }
        }


    def import_sections_articles(self):

        wizard = self.env.ref(
            'e2yun_agreement_docx_import.view_action_agreement_import_doc')
        return {
            'name': '导入章节与条款',
            'type': 'ir.actions.act_window',
            'view_type': 'form',
            'view_mode': 'form',
            'res_model': 'agreement.download.doc',
            'views': [(wizard.id, 'form')],
            'view_id': wizard.id,
            'target': 'new',
            'context': {
                'active_id': self.id,
                'import_type':'import_sections_articles'
            }

        }

    def import_appendix(self):
        wizard = self.env.ref(
            'e2yun_agreement_docx_import.view_action_agreement_import_doc')
        return {
            'name': '导入附录',
            'type': 'ir.actions.act_window',
            'view_type': 'form',
            'view_mode': 'form',
            'res_model': 'agreement.download.doc',
            'views': [(wizard.id, 'form')],
            'view_id': wizard.id,
            'target': 'new',
            'context': {
                'active_id': self.id,
                'import_type':'import_appendix'
            }

        }


class AgreementWordData(models.Model):  #条款
    _name = "agreement.word.data"
    master_word_id = fields.Integer('Master Word Id')
    the_editor = fields.Boolean('The Editor')
    sequence = fields.Integer(string="Sequence")
    agreement_id = fields.Integer()
    content = fields.Text(string="Clause Content")
    old_text=fields.Text()
    new_text = fields.Text()
    detail=fields.Boolean()
    edit_type=fields.Char()
    alignment=fields.Char()
    font_name=fields.Char()
    font_size=fields.Char()




class AgreementRecital(models.Model):  #概述
    _inherit = "agreement.recital"
    edit_type=fields.Char()
    alignment=fields.Char()
    font_name=fields.Char()
    font_size=fields.Char()
    master_word_id = fields.Integer('Master Word Id')
    agreement_placeholder_id=fields.Many2one('agreement.placeholder',string="占位符")
    def write(self, vals):
        if 'content' in vals.keys():
            import re
            str=vals['content']
            n = re.findall(r"a*{(.+?)}", str)  #
            if len(n) > 0:
                content = ""
            else:
                content = str
            strTemp = str
            agreement_placeholder_obj= self.env['agreement.placeholder']  # wordData
            for i, data in enumerate(n):
                strsTemp = strTemp.split("${" + data + "}")
                querName="${" + data + "}"
                agreement_placeholder=agreement_placeholder_obj.search([('name' ,'=',querName)])
                if strsTemp[0]:
                    content = content + strsTemp[0] + agreement_placeholder.text
                if strsTemp[1]:
                    strTemp = strsTemp[1]
                    if i == len(n) - 1:
                        content = content + strTemp
            vals['content']=content
        return super(AgreementRecital,self).write(vals)

class AgreementSection(models.Model):  #章节-条款
    _inherit = "agreement.section"
    edit_type = fields.Char()
    alignment = fields.Char()
    font_name = fields.Char()
    font_size = fields.Char()
    master_word_id = fields.Integer('Master Word Id')


class AgreementAppendix(models.Model):  #附录
    _inherit = "agreement.appendix"
    edit_type = fields.Char()
    alignment = fields.Char()
    font_name = fields.Char()
    font_size = fields.Char()
    master_word_id = fields.Integer('Master Word Id')






